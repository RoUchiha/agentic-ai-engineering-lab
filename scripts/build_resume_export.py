from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "resume" / "Agentic_AI_Resume_Integration_Guide.docx"
NAVY = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GRAY = RGBColor(92, 101, 115)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
RED = RGBColor(155, 28, 28)


def set_font(run, name="Calibri", size=11, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def set_cell_like_padding(paragraph, before=8, after=8, left=10, right=10):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.left_indent = Pt(left)
    paragraph.paragraph_format.right_indent = Pt(right)


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, separate, text, end])
    set_font(run, size=9, color=GRAY)


def add_bullet(doc, text, bold_lead=None):
    p = doc.add_paragraph(style="Portfolio Bullet")
    p.paragraph_format.keep_together = True
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_font(lead, bold=True)
        body = p.add_run(text[len(bold_lead):])
        set_font(body)
    else:
        run = p.add_run(text)
        set_font(run)
    return p


def add_numbered(doc, number, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    marker = p.add_run(f"{number}. ")
    set_font(marker, color=BLUE, bold=True)
    run = p.add_run(text)
    set_font(run)
    return p


def add_label_paragraph(doc, label, text, fill=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.18
    if fill:
        shade(p, fill)
        set_cell_like_padding(p)
    label_run = p.add_run(f"{label}: ")
    set_font(label_run, color=DARK_BLUE, bold=True)
    value_run = p.add_run(text)
    set_font(value_run)
    return p


def add_star(doc, title, situation, task, action, result):
    doc.add_heading(title, level=2)
    add_label_paragraph(doc, "Situation", situation)
    add_label_paragraph(doc, "Task", task)
    add_label_paragraph(doc, "Action", action)
    add_label_paragraph(doc, "Result", result, fill=LIGHT_GRAY)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    bullet_style = doc.styles.add_style("Portfolio Bullet", WD_STYLE_TYPE.PARAGRAPH)
    bullet_style.base_style = normal
    bullet_style.paragraph_format.left_indent = Inches(0.5)
    bullet_style.paragraph_format.first_line_indent = Inches(-0.25)
    bullet_style.paragraph_format.space_after = Pt(6)
    bullet_style.paragraph_format.line_spacing = 1.25
    bullet_style.element.get_or_add_pPr().append(_numbering_properties(1))

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    run = header.add_run("AGENTIC AI  |  RESUME INTEGRATION GUIDE")
    set_font(run, size=8.5, color=GRAY, bold=True)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_before = Pt(0)
    lead = footer.add_run("Roshaan Singh  |  Page ")
    set_font(lead, size=9, color=GRAY)
    add_field(footer, "PAGE")


def _numbering_properties(num_id):
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    return num_pr


def build():
    doc = Document()
    configure_document(doc)
    props = doc.core_properties
    props.title = "Agentic AI Resume Integration Guide"
    props.subject = "Truthful resume and interview examples for agentic AI engineering"
    props.author = "Roshaan Singh"
    props.keywords = "agentic AI, resume, Citi, portfolio, engineering"

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_before = Pt(42)
    kicker.paragraph_format.space_after = Pt(14)
    run = kicker.add_run("CAREER POSITIONING GUIDE  /  2026")
    set_font(run, size=9.5, color=BLUE, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run("Agentic AI Resume\nIntegration Guide")
    set_font(run, size=28, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(28)
    run = subtitle.add_run("Believable financial-services examples grounded in an inspectable engineering portfolio")
    set_font(run, size=12.5, color=GRAY, italic=True)

    callout = doc.add_paragraph()
    shade(callout, LIGHT_BLUE)
    set_cell_like_padding(callout, before=12, after=12, left=16, right=16)
    callout.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lead = callout.add_run("Accuracy first. ")
    set_font(lead, size=10.5, color=NAVY, bold=True)
    body = callout.add_run("Use only examples that match work you actually performed. Replace bracketed metrics with values you can defend, and never imply that the public portfolio uses Citi systems or data.")
    set_font(body, size=10.5, color=NAVY)

    intro = doc.add_paragraph()
    intro.paragraph_format.space_before = Pt(26)
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = intro.add_run("Prepared for Roshaan Singh  |  Personal portfolio companion")
    set_font(run, size=9.5, color=GRAY)

    doc.add_page_break()
    doc.add_heading("1. Recommended skills section", level=1)
    add_label_paragraph(doc, "Agentic AI", "LLM workflow orchestration, tool and function calling, retrieval-augmented generation (RAG), human-in-the-loop approvals, structured outputs, prompt and workflow evaluation, guardrails, model observability, and cost/latency optimization.")
    add_label_paragraph(doc, "Engineering", "Python, FastAPI, Pydantic, TypeScript, React, Next.js, REST APIs, SQL, Docker, GitHub Actions, pytest, Vitest, Prometheus, PromQL, Grafana, and OpenTelemetry.")
    add_label_paragraph(doc, "Enterprise AI", "AI risk controls, least-privilege tool design, audit logging, data privacy, change management, failure recovery, idempotent workflows, and stakeholder communication.")

    doc.add_heading("2. Citi experience bullets - conservative", level=1)
    conservative = [
        "Prototyped an AI-assisted operations workflow that gathered case evidence from approved sources, produced a structured recommendation, and routed high-impact actions for human approval rather than autonomous execution.",
        "Built retrieval and citation checks for internal procedure content, helping users trace AI-generated answers back to relevant policy sections and reducing unsupported responses during testing.",
        "Added repeatable evaluation cases for task completion, evidence coverage, and policy compliance, enabling workflow changes to be regression-tested before demonstration or release.",
        "Designed typed Python API contracts and validation rules for AI workflow inputs and outputs, improving error handling and making downstream integration behavior easier to test.",
        "Instrumented prototype agent runs with trace IDs, tool usage, latency, retry, and estimated token-cost metrics to support troubleshooting and engineering trade-off discussions.",
        "Built Grafana dashboards for workflow SLOs, model cost versus evaluation quality, and tool reliability using bounded Prometheus metrics and version-controlled PromQL queries.",
        "Defined prototype alerts for low workflow success, high p95 latency, and elevated tool failure rates, with demonstration thresholds pending production baselines.",
        "Documented agent boundaries, approval points, failure modes, and safe defaults in plain language for engineering, operations, and risk stakeholders.",
    ]
    for item in conservative:
        add_bullet(doc, item)

    doc.add_heading("3. Metric-based bullets - use only with real measurements", level=1)
    metric_bullets = [
        "Reduced average investigation preparation time by [X%] in a controlled pilot by using an agent workflow to assemble approved evidence and draft a reviewer-ready case summary.",
        "Improved citation coverage from [baseline%] to [result%] across [N] policy questions by adding version-aware retrieval and automated evidence checks.",
        "Cut manual handoffs by [N steps/case] for [workflow] by generating structured work plans while preserving supervisor approval for customer or financial actions.",
        "Increased evaluation pass rate to [X%] across [N] golden cases by introducing regression tests for groundedness, task success, and policy compliance.",
        "Lowered prototype inference cost by [X%] through model routing, context trimming, caching, or deterministic handling of low-complexity steps while maintaining the agreed quality threshold.",
    ]
    for item in metric_bullets:
        add_bullet(doc, item)

    doc.add_heading("4. Personal project entry", level=1)
    project = doc.add_paragraph()
    project.paragraph_format.space_after = Pt(8)
    r = project.add_run("Agentic Systems Lab")
    set_font(r, size=12, color=NAVY, bold=True)
    r = project.add_run("  |  TypeScript, Next.js, Python, FastAPI, Prometheus, Grafana, Docker")
    set_font(r, size=10.5, color=GRAY)
    for item in [
        "Built a runnable enterprise-agent reference platform featuring multi-role orchestration, typed tool contracts, RAG evidence, policy-as-code, human approval, evaluation gates, and run-level observability.",
        "Implemented three synthetic financial-services workflows - payment exception triage, policy evidence Q&A, and customer remediation planning - with deterministic local execution and no required API key.",
        "Added TypeScript and Python test suites, golden evaluation cases, containerized services, CI quality gates, architecture documentation, and an optional OpenAI Responses API adapter.",
        "Added a local Prometheus and Grafana lab with dashboards-as-code, alert provisioning, synthetic traffic, and CI validation for metrics and dashboard configuration.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("5. LinkedIn or profile summary", level=1)
    profile = doc.add_paragraph()
    shade(profile, LIGHT_GRAY)
    set_cell_like_padding(profile, before=12, after=12, left=16, right=16)
    run = profile.add_run("Agentic AI engineer focused on turning LLM capabilities into controlled, testable enterprise workflows. I build retrieval and tool-using systems with explicit permissions, human approval for high-impact actions, automated evaluations, and production-minded observability. My background in financial-services technology helps me translate operations and risk requirements into reliable Python and TypeScript services that stakeholders can inspect and trust.")
    set_font(run, size=10.5, color=NAVY)

    doc.add_page_break()
    doc.add_heading("6. Truthful STAR interview examples", level=1)
    lead = doc.add_paragraph("Keep each story anchored to a real artifact, decision, and stakeholder. If a result was a prototype insight rather than a production KPI, say that clearly.")
    lead.paragraph_format.space_after = Pt(10)

    add_star(doc, "Operations investigation assistant", "Analysts had to collect information from several approved sources before reviewing an exception.", "Explore whether AI could reduce preparation work without allowing it to make the final high-impact decision.", "Built a bounded workflow that retrieved evidence, retained source identifiers, drafted a recommendation, and sent write actions to a maker-checker approval queue. Added tests for unknown tools, missing evidence, and approval behavior.", "Demonstrated a repeatable reviewer-ready case package and established the controls and metrics needed for a measured pilot. Add a quantified result only if it was actually measured.")
    add_star(doc, "Policy-grounded assistant", "Employees needed faster answers from long procedure documents, but unsupported answers were unacceptable.", "Make answers traceable and measurable.", "Split approved documents into versioned evidence units, retrieved the most relevant passages, required source IDs in the response, and scored citation coverage and groundedness in a golden test set.", "Produced answers reviewers could verify quickly and made unsupported or stale responses visible during testing.")
    add_star(doc, "AI engineering quality gates", "Prompt and model changes could improve one example while silently degrading another.", "Create a release discipline for agent behavior.", "Added scenario-level expected outcomes, minimum evaluation thresholds, policy unit tests, and CI checks for both TypeScript and Python services.", "Turned subjective demo review into a repeatable pass/fail signal and made regressions easier to diagnose from traces and metric breakdowns.")
    add_star(doc, "Agent observability and SLOs", "Agent failures could originate in workflow logic, model behavior, or downstream tools, making a single health check insufficient.", "Create operational views that separated user-facing reliability, cost-quality trade-offs, and dependency health.", "Instrumented a typed API with bounded Prometheus metrics, wrote PromQL for three provisioned Grafana dashboards, and defined alerts for success rate, p95 latency, and tool errors.", "Produced a repeatable local observability lab and a concrete baseline for choosing production SLOs and alert thresholds after real traffic became available.")

    doc.add_heading("7. Claims to avoid", level=1)
    for item in [
        "Deployed autonomous agents at Citi - unless you truly did and can discuss the scope and controls.",
        "Reduced costs by millions - without an approved measurement and attribution method.",
        "Eliminated hallucinations or achieved 100% accuracy - neither claim is credible.",
        "Built Citi's AI platform - if you built a component, prototype, or contributed as part of a team.",
        "Any confidential system names, customer data, internal control identifiers, or non-public model decisions.",
    ]:
        p = add_bullet(doc, item)
        for run in p.runs:
            run.font.color.rgb = RED

    doc.add_heading("8. Final tailoring checklist", level=1)
    checklist = [
        "Match each bullet to a real task, artifact, stakeholder, and outcome you can explain.",
        "Replace generic words with the actual workflow category, while removing confidential names.",
        "Add numbers only when you know the baseline, measurement window, and sample size.",
        "Separate prototype, pilot, and production clearly.",
        "Keep the portfolio listed as a personal project unless it was formally part of your Citi role.",
    ]
    for number, item in enumerate(checklist, start=1):
        add_numbered(doc, number, item)

    closing = doc.add_paragraph()
    shade(closing, LIGHT_BLUE)
    set_cell_like_padding(closing, before=12, after=12, left=16, right=16)
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = closing.add_run("Best positioning: bounded autonomy, measurable quality, and enterprise-ready controls.")
    set_font(run, size=11, color=NAVY, bold=True)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
